import PyPDF2
import docx

# # Open pdf file
# pdfFileObj = open('_PN-83_B-03430_doc.pdf', 'rb')
# pdfReader = PyPDF2.PdfReader(pdfFileObj)
# len(pdfReader.pages)

# pageObj = pdfReader.pages[1]
# pageObj.extract_text()
# pdfFileObj.close()

# # Copying pages

# import PyPDF2
# pdf1File = open('meetingminutes.pdf', 'rb')
# pdf2File = open('meetingminutes2.pdf', 'rb')
# pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
# pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
# pdfWriter = PyPDF2.PdfFileWriter()

# for pageNum in range(pdf1Reader.numPages):
#         pageObj = pdf1Reader.getPage(pageNum)
#         pdfWriter.addPage(pageObj)

# for pageNum in range(pdf2Reader.numPages):
#         pageObj = pdf2Reader.getPage(pageNum)
#         pdfWriter.addPage(pageObj)

# pdfOutputFile = open('combinedminutes.pdf', 'wb')
# pdfWriter.write(pdfOutputFile)
# pdfOutputFile.close()
# pdf1File.close()
# pdf2File.close()

# # Open doxc

# doc = docx.Document('demo.docx')
# len(doc.paragraphs)
# doc.paragraphs[0].text
# len(doc.paragraphs[0].runs) # Show different styles of text???
# doc.paragraphs[0].runs[0].text

# # Get full text

# def getText(filename):
#     doc = docx.Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         fullText.append(para.text)
#     return '\n'.join(fullText)

# print(getText('demo.docx'))

# import readDocx
# print(readDocx.getText('demo.docx'))

# # Write to word

# doc = docx.Document()
# doc.add_paragraph('Hello, world!')
# doc.add_heading('Header 0', 0)
# doc.add_picture('zophie.png', width=docx.shared.Inches(1),
# height=docx.shared.Cm(4))
# doc.save('helloworld.docx')

# Creating PDF
# Working only on win with word

import win32com.client # install with "pip install pywin32==224"

wordFilename = 'C:/Users/pawel/Documents/GitHub/Learning-python/demo.docx'
pdfFilename = 'C:/Users/pawel/Documents/GitHub/Learning-python/demo.pdf'

doc = docx.Document()
# Code to create Word document goes here.
doc.add_paragraph('Hello, world!')
doc.add_heading('Header 0', 0)
doc.save(wordFilename)


wdFormatPDF = 17 # Word's numeric code for PDFs.
wordObj = win32com.client.Dispatch('Word.Application')

docObj = wordObj.Documents.Open(wordFilename)
docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
docObj.Close()
wordObj.Quit()